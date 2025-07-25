"""
=============================================================================
AIVS Accounting RAG API
=============================================================================

Version: 2.0.0
Last Updated: 2025-04-27
Author: Michael Williams

-----------------------------------------------------------------------------

DESCRIPTION:
This Flask-based API provides Accounting and Bookkeeping professionals
with AI-generated structured responses based on UK accounting standards.
It uses a FAISS vector database built from accounting, bookkeeping,
corporate governance, and financial compliance source materials.

Each query:
- Is embedded and searched against the Accounting FAISS index.
- Retrieves relevant context chunks.
- Sends the combined context and question to OpenAI GPT-4.
- Generates a structured reply including:
    • Client-facing response
    • Action recommendations
    • Relevant UK accounting standards references
- Outputs a professionally formatted Word document (.docx).
- Sends the resulting document to the User, Supervisor, and HR via Mailjet.

-----------------------------------------------------------------------------

KEY COMPLIANCE STANDARDS:
- UK GAAP (Generally Accepted Accounting Practice)
- IFRS (International Financial Reporting Standards)
- FRS 102 (Small entities)
- FRS 105 (Micro-entities)
- UK Companies Act 2006
- UK Corporate Governance Code
- UK Financial Practices Code
- Insolvency Act 1986
- HMRC Reporting and Compliance Requirements

-----------------------------------------------------------------------------

TECHNOLOGY STACK:
- Python 3.9+
- Flask
- FAISS (CPU version)
- OpenAI API (GPT-4, Embedding 3-small)
- Mailjet Email Delivery (via Postmarker)
- Python-Docx (Word document generation)
- Render or Local Hosting (Gunicorn supported)

-----------------------------------------------------------------------------

NOTES:
- This version is Accounting and Bookkeeping only.
- No Police, Strategic Management, or other disciplines are included.
- Old discipline pathways (e.g., Police Procedure) have been redacted but preserved for future use.

=============================================================================
"""
import os
import faiss
import pickle
import json
import base64
import datetime
import re
import numpy as np
import requests
import textwrap
import openai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from zoneinfo import ZoneInfo  # Python 3.9+ 


__version__ = "v1.0.7-test"
print(f"🚀 API Version: {__version__}")

def add_markdown_bold(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
print("🔒 OPENAI_API_KEY exists?", bool(OPENAI_API_KEY))
client = OpenAI(api_key=OPENAI_API_KEY)

app = Flask(__name__)
CORS(app, origins=["https://www.aivs.uk"])

@app.route("/", methods=["GET"])
def home():
    return "✅ Business API is running", 200

@app.after_request
def apply_cors_headers(response):
    response.headers.add("Access-Control-Allow-Origin", "https://www.aivs.uk")
    response.headers.add("Access-Control-Allow-Headers", "Content-Type")
    response.headers.add("Access-Control-Allow-Methods", "POST, OPTIONS")
    return response

@app.route("/ping", methods=["POST", "OPTIONS"])
def ping():
    if request.method == "OPTIONS":
        return '', 204
    return jsonify({"message": "pong"})

# Load FAISS index
try:
    faiss_index = faiss.read_index("data/accounting/accounting_chunks.index")
    with open("data/accounting/metadata.pkl", "rb") as f:
        metadata = pickle.load(f)
    print("✅ Accounting FAISS index and metadata loaded.")
except Exception as e:
    faiss_index = None
    metadata = []
    print("⚠️ Failed to load Accounting FAISS index:", str(e))

def ask_gpt_with_context(data, context):
    query = data.get("query", "")
    job_title = data.get("job_title", "Not specified")
    seniority_level = data.get("seniority_level", "Not specified") 
    timeline = data.get("timeline", "Not specified")
    discipline = data.get("discipline", "Not specified")
    site = data.get("site", "Not specified")
    funnel_1 = data.get("funnel_1", "Not specified")
    funnel_2 = data.get("funnel_2", "Not specified")
    funnel_3 = data.get("funnel_3", "Not specified")

    prompt = f"""
You are responding to a professional business query via a secure reporting system.

All responses must:
- Be based on correct UK financial standards, accounting regulations, business risk practices, or strategic management theory.
- Use British English spelling and tone.

### Enquiry:
\"{query}\"

### Context from FAISS Index:
{context}


### Additional Focus:
- Support Need: {funnel_1}
- Current Status: {funnel_2}
- Follow-Up Expectation: {funnel_3}

### Your Task:
Please generate a structured response that includes:

1. **Client Reply** – plain English appropriate for senior business audiences.
2. **Action Sheet** – bullet-point recommended actions.
3. **Policy or Standard Notes** – cite relevant accounting standards, regulatory codes, or strategic management principles if applicable.
"""
    return generate_reviewed_response(prompt,discipline)

def generate_reviewed_response(prompt,discipline,):
    print("📢 Sending initial GPT prompt...")

    completion = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=1800  # Allow longer initial output
    )
    initial_response = completion.choices[0].message.content.strip()

    # 📏 Initial GPT response length
    print(f"📏 Initial GPT response length: {len(initial_response)} characters")

    # ⛔ Skip review if too big
    if len(initial_response) > 1500:
        print("⚡ Skipping review — using initial GPT response directly.")
        return initial_response

    # 🔄 Otherwise review normally
    print("🔄 Reviewing GPT response...")

    # 🧼 Strip polite sign-offs
    initial_response = re.sub(
        r'(Best regards,|Yours sincerely,|Kind regards,)[\s\S]*$',
        '',
        initial_response,
        flags=re.IGNORECASE
    ).strip()

    # ✂️ Trim FAISS context and limit input length
    stripped_response = initial_response.split("### Context from FAISS Index:")[0].strip()
    stripped_response = stripped_response[:2000]  # Safe upper limit

    # 🧠 Build review prompt using textwrap.dedent
    if discipline == "Accounting":
          review_prompt = textwrap.dedent(f"""\
       You are acting as a UK Chartered Accountant.
                                          
    Review and formally rewrite the following draft according to:
    - UK GAAP
    - UK Accounting Standards
    - UK Companies Act 2006
    - UK Corporate Governance Code
    - Other relevant UK company law and governance standards
    - UK Financial Practices Code                
    - IFRS
    - Insolvency Act 1986
    - Other relevant UK accounting regulations
    Maintain strict professional tone, British English spelling, and correct accounting terminology.                        
       --- START RESPONSE ---
       {stripped_response}
       --- END RESPONSE ---
       """)
    #elif discipline == "Police Procedure":
          #review_prompt = textwrap.dedent(f"""\
        #You are acting as a UK police Professional Standards Officer or Custody Sergeant.

        #Rewrite the following draft as formal, clear procedural guidance for frontline officers. Focus on correct application of UK law (PACE, Criminal Law Act, etc.), internal policies, and officer conduct.

        #Structure as:
        #- ISSUE SUMMARY
        #- APPLICABLE LAW
        #- PROCEDURAL GUIDANCE
        #- RISK NOTES

        #No command tone needed. Use neutral, professional legal explanation style.

        #--- START RESPONSE ---
        #{stripped_response}
        #--- END RESPONSE ---
        #""")

    else:
        review_prompt = textwrap.dedent(f"""\
        Please clean and improve the following structured response while maintaining professional tone and factual accuracy.

        --- START RESPONSE ---
        {stripped_response}
        --- END RESPONSE ---
        """)

    # 🚀 Request GPT review with tight limits
    review_completion = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": review_prompt}],
        temperature=0,
        max_tokens=700,  # Trimmed to avoid Render crashes
        timeout=15        # Optional cap to prevent long hangs
    )

    reviewed_response = review_completion.choices[0].message.content.strip()
    print(f"✅ Reviewed response length: {len(reviewed_response)} characters")
    return reviewed_response

def send_email_mailjet(to_emails, subject, body_text, attachments=[], full_name=None, supervisor_name=None):
    MAILJET_API_KEY = os.getenv("MJ_APIKEY_PUBLIC")
    MAILJET_SECRET_KEY = os.getenv("MJ_APIKEY_PRIVATE")

    messages = []

    for recipient in to_emails:
        role = recipient["Name"]
        email = recipient["Email"]

        # Customise the message per role
        if role == full_name:
            text_body = f"""To: {full_name},

Please find attached the AI-generated analysis based on your query submitted on {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")}.
"""
        elif role == supervisor_name:
            text_body = f"""To: {supervisor_name},

Please review the attached report submitted by {full_name}. It contains AI-generated analysis for internal review.
"""
        elif role == "HR Department":
            text_body = f"""To: HR Department,

This document was generated following a query submitted by {full_name}. Please file or follow up according to internal procedures.
"""
        else:
            text_body = f"Attached is an AI-generated analysis regarding {full_name}."

        messages.append({
            "From": {
                "Email": "noreply@securemaildrop.uk",
                "Name": "Secure Maildrop"
            },
            "To": [{"Email": email, "Name": role}],
            "Subject": subject,
            "TextPart": text_body,
            "HTMLPart": f"<pre>{text_body}</pre>",
            "Attachments": [
                {
                    "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "Filename": os.path.basename(file_path),
                    "Base64Content": base64.b64encode(open(file_path, "rb").read()).decode()
                }
                for file_path in attachments
            ]
        })

    response = requests.post(
        "https://api.mailjet.com/v3.1/send",
        auth=(MAILJET_API_KEY, MAILJET_SECRET_KEY),
        json={"Messages": messages}
    )

    print(f"📤 Mailjet status: {response.status_code}")
    print(response.json())
    return response.status_code, response.json()

@app.route("/generate", methods=["POST"])
def generate_response():
    print("📥 /generate route hit")
    try:
        data = request.get_json()
        print("🔎 Payload received:", data)
    except Exception as e:
        print("❌ Error parsing JSON:", e)
        return jsonify({"error": "Invalid JSON input"}), 400

    query_text = data.get("query")
    full_name = data.get("full_name", "User")
    user_email = data.get("user_email")
    supervisor_email = data.get("supervisor_email")
    hr_email = data.get("hr_email")
    supervisor_name = data.get("supervisor_name", "Supervisor")
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")

    if faiss_index:
        query_vector = client.embeddings.create(
            input=[query_text.replace("\n", " ")],
            model="text-embedding-3-small"
        ).data[0].embedding

        D, I = faiss_index.search(np.array([query_vector]).astype("float32"), 2)

        matched_chunks = []
        for i in I[0]:
            chunk_file = metadata[i]["chunk_file"]
            with open(f"data/{chunk_file}", "r", encoding="utf-8") as f:
                matched_chunks.append(f.read().strip())

        context = "\n\n---\n\n".join(matched_chunks)

        # Redact sensitive info
        #sensitive_names = ["Wiltshire Police", "Humberside Police", "Avon and Somerset Police"]
        #for name in sensitive_names:
            #context = context.replace(name, "the relevant police force")

        #context = re.sub(r'\b(PC|SGT|CID)?\d{3,5}\b', '[badge number]', context, flags=re.IGNORECASE)

    else:
        context = "Policy lookup not available (FAISS index not loaded)."

    answer = ask_gpt_with_context(data, context)

    # ✅ Remove repeated '### ORIGINAL QUERY' section if GPT included it
    answer = re.sub(r"### ORIGINAL QUERY\s*[\r\n]+.*?(?=###|\Z)", "", answer, flags=re.IGNORECASE | re.DOTALL).strip()
    
    print(f"🧠 GPT answer: {answer[:80]}...")
    
    discipline = data.get("discipline", "Not specified")
    discipline_folder = discipline.lower().replace(" ", "_")
    output_path = f"output/{discipline_folder}"
    os.makedirs(output_path, exist_ok=True)

    doc_path = f"{output_path}/{full_name.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    
    doc = Document()

    # ✅ Apply default document style
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # Document styling
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"RESPONSE FOR {full_name.upper()}")
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # ✅ UK-style timestamp
    uk_time = datetime.now(ZoneInfo("Europe/London"))
    generated_datetime = uk_time.strftime("%d %B %Y at %H:%M:%S (%Z)")
    doc.add_paragraph(f"Generated: {generated_datetime}")

    # 🔹 ORIGINAL QUERY heading
    para_query_heading = doc.add_paragraph()
    run_heading = para_query_heading.add_run("ORIGINAL QUERY")
    run_heading.bold = True
    run_heading.font.name = 'Arial'
    run_heading.font.size = Pt(11)
    run_heading.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Divider ABOVE the query text
    divider_above = doc.add_paragraph()
    divider_above_run = divider_above.add_run("────────────────────────────────────────────")
    divider_above_run.font.name = 'Arial'
    divider_above_run.font.size = Pt(10)
    divider_above_run.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Italicised query
    para_query_text = doc.add_paragraph()
    run_query = para_query_text.add_run(f'"{query_text.strip()}"')
    run_query.italic = True
    run_query.font.name = 'Arial'
    run_query.font.size = Pt(11)
    run_query.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Divider BELOW the query text
    divider_below = doc.add_paragraph()
    divider_below_run = divider_below.add_run("────────────────────────────────────────────")
    divider_below_run.font.name = 'Arial'
    divider_below_run.font.size = Pt(10)
    divider_below_run.font.color.rgb = RGBColor(0, 0, 0)

    # ✅ Bold header: "AI RESPONSE"
    para1 = doc.add_paragraph()
    run1 = para1.add_run("AI RESPONSE")
    run1.bold = True
    run1.font.name = 'Arial'
    run1.font.size = Pt(11)
    run1.font.color.rgb = RGBColor(0, 0, 0)

    # ✅ Subheader: "Note: ..."
    para2 = doc.add_paragraph()
    run2 = para2.add_run("Note: This report was prepared using AI analysis based on the submitted query.")
    run2.bold = True
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0, 0, 0)

    # ✅ ORIGINAL QUERY label✅ ORIGINAL QUERY label
    # ✅ ORIGINAL QUERY labelpara3 = doc.add_paragraph()
    # ✅ ORIGINAL QUERY labelrun3 = para3.add_run("ORIGINAL QUERY:\n")
    # ✅ ORIGINAL QUERY labelrun3.bold = True
    # ✅ ORIGINAL QUERY labelrun3.font.name = 'Arial'
    # ✅ ORIGINAL QUERY labelrun3.font.size = Pt(11)
    # ✅ ORIGINAL QUERY labelrun3.font.color.rgb = RGBColor(0, 0, 0)
 #
    # 🔹 ORIGINAL QUERY heading
    #para_query_heading = doc.add_paragraph()
    #run_heading = para_query_heading.add_run("ORIGINAL QUERY")
    #run_heading.bold = True
    #run_heading.font.name = 'Arial'
    #run_heading.font.size = Pt(11)
    #run_heading.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Divider ABOVE the query text
    # divider_above = doc.add_paragraph()
     #divider_above_run = divider_above.add_run("────────────────────────────────────────────")
     #divider_above_run.font.name = 'Arial'
     #divider_above_run.font.size = Pt(10)
     #divider_above_run.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Italicised query
    #para_query_text = doc.add_paragraph()
    #run_query = para_query_text.add_run(f'"{query_text.strip()}"')
    #run_query.italic = True
    #run_query.font.name = 'Arial'
    #run_query.font.size = Pt(11)
    #run_query.font.color.rgb = RGBColor(0, 0, 0)

    # 🔹 Divider BELOW the query text
    divider_below = doc.add_paragraph()
    divider_below_run = divider_below.add_run("────────────────────────────────────────────")
    divider_below_run.font.name = 'Arial'
    divider_below_run.font.size = Pt(10)
    divider_below_run.font.color.rgb = RGBColor(0, 0, 0)
    # ✅ Actual query content
    #doc.add_paragraph(query_text.strip())

    #divider = doc.add_paragraph()
    #divider.add_run("────────────────────────────────────────────").font.size = Pt(10)

    # Split answer into structured sections
    sections = re.split(r'^### (.*?)\n', answer, flags=re.MULTILINE)
    structured = {}
    current_title = None

    for i, part in enumerate(sections):
        content = part.strip()
        if i == 0 and content:
            content = re.sub(r'^\s*Enquirer Reply\s*', '', content, flags=re.IGNORECASE)
            content = re.sub(r'^\s*Hello,\s*', '', content, flags=re.IGNORECASE)
            structured["Enquirer Reply"] = content
        elif i % 2 == 1:
            current_title = content
        elif i % 2 == 0 and current_title:
            if current_title.lower() in ["enquirer reply", "initial response"]:
                lines = content.splitlines()
                cleaned_lines = [line for line in lines if not re.match(r'^\s*(enquirer reply|hello,?)\s*$', line, flags=re.IGNORECASE)]
                content = "\n".join(cleaned_lines).strip()
            structured[current_title] = content
            # ✅ Added debug line here 16 38
            print("🔍 Structured keys:", list(structured.keys()))

    if not structured:
        print("⚠️ GPT returned unstructured content. Using entire answer as 'Initial Response'.")
        structured["Initial Response"] = answer.strip()

    rename = {"Enquirer Reply": "Initial Response"}
    for title in structured:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(rename.get(title, title).upper())
        heading_run.bold = True
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)

        if title in ["Action Sheet", "Policy Notes"]:
           lines = structured[title].splitlines()
           for line in lines:
              clean = line.strip()
              clean = re.sub(r'^[-•–]?\s*\d+[.)]?\s*', '', clean)
              clean = re.sub(r'^[-•–]\s*', '', clean)
              if clean:
                  para = doc.add_paragraph(clean, style='List Number')
                  para.paragraph_format.left_indent = Mm(5) 

         #elif title == "Policy Notes":
            #lines = structured[title].splitlines()
            #for line in lines:
               #clean = line.strip()
               #clean = re.sub(r'^[-•–]?\s*\d+[.)]?\s*', '', clean)
               #clean = re.sub(r'^[-•–]\s*', '', clean)
               #if clean:
                   #para = doc.add_paragraph(clean, style='List Number')
                   #para.paragraph_format.left_indent = Mm(5)         
        else:
            # Default fallback for all other sections
            cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', structured[title])
            para = doc.add_paragraph()
            run = para.add_run(cleaned_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    doc.add_paragraph("────────────────────────────────────────────")
    doc.add_paragraph("This document was generated by AIVS Software Limited using AI assistance (OpenAI). Please review for accuracy and relevance before taking any formal action.")
    doc.add_paragraph("© AIVS Software Limited 2025. All rights reserved.")
    doc.add_paragraph(datetime.now(ZoneInfo("Europe/London")).strftime("Report generated on %d %B %Y at %H:%M:%S (%Z)"))

    doc.save(doc_path)
    print(f"📄 Word saved: {doc_path}")

    recipients = []
    if user_email:
        recipients.append({"Email": user_email, "Name": full_name})
    if supervisor_email:
        recipients.append({"Email": supervisor_email, "Name": supervisor_name})
    if hr_email:
        recipients.append({"Email": hr_email, "Name": "HR Department"})

    if not recipients:
        return jsonify({"error": "No valid email addresses provided."}), 400

    subject = f"AI Analysis for {full_name} - {timestamp}"
    body_text = f"""To: {full_name},

Please find attached the AI-generated analysis based on your query submitted on {timestamp}.
"""

    status, response = send_email_mailjet(
        to_emails=recipients,
        subject=subject,
        body_text=body_text,
        attachments=[doc_path],
        full_name=full_name,
        supervisor_name=supervisor_name
    )

    return jsonify({
        "status": "ok",
        "message": "✅ OpenAI-powered response generated, AI reviewed and email successfully sent.",
        "disclaimer": "This document was generated by AIVS Software Limited using AI assistance (OpenAI). Please review for accuracy and relevance before taking any formal action.",
        "copyright": "© AIVS Software Limited 2025. All rights reserved.",
        "context_preview": context[:200],
        "mailjet_status": status,
        "mailjet_response": response
    })
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
